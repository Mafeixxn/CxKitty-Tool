#!/usr/bin/env python3
"""
学习通题目爬虫 - 基于 CxKitty API 层
功能: 登录 → 拉取课程 → 拉取章节 → 提取题目 → 导出 Word (.docx)

用法:
    python scraper.py                    # 交互式选择课程 (默认直连)
    python scraper.py --all              # 爬取所有课程
    python scraper.py --course "课程名"   # 爬取指定课程
    python scraper.py --proxy 10808      # 通过代理端口访问
"""

import os
import sys
import json
import time
import argparse
import traceback
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 确保能导入同目录下的 cxapi 模块
sys.path.insert(0, str(Path(__file__).parent))

from cxapi.api import ChaoXingAPI


# ============================================================
# Windows 密码输入 (msvcrt, 兼容所有终端)
# ============================================================
def _read_password(prompt: str = "密码: ") -> str:
    """跨平台密码输入, Windows 下用 msvvcrt 逐字符读取"""
    if sys.platform == "win32":
        import msvcrt
        sys.stdout.write(prompt)
        sys.stdout.flush()
        chars = []
        while True:
            ch = msvcrt.getch()
            # Ctrl+C / Ctrl+Z
            if ch in (b"\x03", b"\x1a"):
                sys.stdout.write("\n")
                raise KeyboardInterrupt
            # Enter
            if ch == b"\r":
                sys.stdout.write("\n")
                sys.stdout.flush()
                break
            # Backspace
            if ch == b"\x08":
                if chars:
                    chars.pop()
                    sys.stdout.write("\b \b")
                    sys.stdout.flush()
                continue
            chars.append(ch.decode("gbk", errors="replace"))
            sys.stdout.write("*")
            sys.stdout.flush()
        return "".join(chars).strip()
    else:
        import getpass
        return getpass.getpass(prompt).strip()

# ============================================================
# 代理配置
# ============================================================
PROXY_HOST = "127.0.0.1"


def setup_proxy(port: int):
    """设置 HTTP 代理环境变量"""
    proxy_url = f"http://{PROXY_HOST}:{port}"
    os.environ["HTTP_PROXY"] = proxy_url
    os.environ["HTTPS_PROXY"] = proxy_url
    os.environ["NO_PROXY"] = "localhost,127.0.0.1"
    print(f"[代理] 已设置代理: {proxy_url}")
    return proxy_url


def clear_proxy():
    """清除代理环境变量"""
    for k in ("HTTP_PROXY", "HTTPS_PROXY", "NO_PROXY"):
        os.environ.pop(k, None)
    print("[代理] 不使用代理, 直连")


def patch_session_proxy(session, proxy_url: str):
    """为已有的 requests.Session 设置代理"""
    session.proxies = {"http": proxy_url, "https": proxy_url}
    session.trust_env = False


# ============================================================
# 输出目录
# ============================================================
OUTPUT_DIR = Path("output")


def sanitize_filename(name: str) -> str:
    """清理文件名中的非法字符"""
    forbidden = '<>:"/\\|?*'
    for ch in forbidden:
        name = name.replace(ch, "_")
    return name.strip()


# ============================================================
# 题目导出
# ============================================================

QUESTION_TYPE_NAMES = {
    0: "单选题",
    1: "多选题",
    2: "填空题",
    3: "判断题",
    4: "简答题",
    5: "名词解释",
    6: "论述题",
    7: "计算题",
    8: "其它",
    9: "分录题",
    10: "资料题",
    11: "连线题",
    13: "排序题",
    14: "完型填空",
    15: "阅读理解",
    18: "口语题",
    19: "听力题",
}


def export_docx(course_name: str, chapters_data: list, output_dir: Path):
    """导出为 Word (.docx) 文件"""
    course_dir = output_dir / sanitize_filename(course_name)
    course_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading(course_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 导出信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}    "
        f"章节数: {len(chapters_data)}    "
        f"题目总数: {sum(len(ch['questions']) for ch in chapters_data)}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    idx = 0
    for ch in chapters_data:
        # 章节标题
        doc.add_heading(f"{ch['label']} {ch['name']}", level=1)

        for q in ch["questions"]:
            idx += 1
            qtype = QUESTION_TYPE_NAMES.get(q["type"], f"未知({q['type']})")

            # 题目
            q_para = doc.add_paragraph()
            q_para.paragraph_format.space_after = Pt(4)
            q_run = q_para.add_run(f"{idx}. ({qtype}) ")
            q_run.bold = True
            q_run.font.size = Pt(11)
            q_para.add_run(q["question"]).font.size = Pt(11)

            # 选项
            if q["options"]:
                for opt_label, opt_text in q["options"].items():
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1)
                    opt_para.paragraph_format.space_after = Pt(2)
                    opt_run = opt_para.add_run(f"{opt_label}. {opt_text}")
                    opt_run.font.size = Pt(10.5)

            # 正确答案
            answer = q.get("answer", "")
            if answer:
                # 尝试将答案字母展开为完整选项文字
                options = q.get("options", {})
                answer_texts = []
                for ch in answer:
                    if ch in options and options[ch]:
                        answer_texts.append(f"{ch}. {options[ch]}")
                    else:
                        answer_texts.append(ch)
                full_answer = "\n".join(answer_texts) if len(answer_texts) > 1 else answer_texts[0]

                ans_para = doc.add_paragraph()
                ans_para.paragraph_format.space_before = Pt(6)
                ans_para.paragraph_format.space_after = Pt(8)
                ans_run = ans_para.add_run(f"✓ 正确答案: {full_answer}")
                ans_run.bold = True
                ans_run.font.size = Pt(10.5)
                ans_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                doc.add_paragraph()  # 空行

    filepath = course_dir / f"{sanitize_filename(course_name)}.docx"
    doc.save(str(filepath))
    print(f"  [Word] 已导出 → {filepath}")


# ============================================================
# 爬虫核心逻辑
# ============================================================


class QuestionScraper:
    """学习通题目爬虫"""

    def __init__(self, proxy_port: int = None, debug: bool = False):
        if proxy_port:
            self.proxy_url = setup_proxy(proxy_port)
        else:
            clear_proxy()
            self.proxy_url = None
        self.api = None
        self.output_dir = OUTPUT_DIR
        self.debug = debug

    def _log_error(self, msg: str, exc: Exception = None):
        """打印错误, debug 模式下输出完整堆栈"""
        print(msg)
        if self.debug and exc:
            traceback.print_exc()

    def login_password(self, phone: str, password: str) -> bool:
        """手机号+密码登录"""
        print("[登录] 正在连接学习通...")
        try:
            self.api = ChaoXingAPI()
        except Exception as e:
            print(f"[错误] 初始化 API 失败: {e}")
            traceback.print_exc()
            return False

        if self.proxy_url:
            patch_session_proxy(self.api.session, self.proxy_url)

        try:
            ok, resp = self.api.login_passwd(phone, password)
        except Exception as e:
            print(f"[错误] 登录请求失败: {e}")
            traceback.print_exc()
            print("[提示] 可能是网络不通, 试试加 --proxy 10808")
            return False

        if ok:
            try:
                self.api.accinfo()
            except Exception as e:
                print(f"[错误] 获取账号信息失败: {e}")
                return False
            print(f"[登录成功] {self.api.acc.name} ({self.api.acc.school})")
            return True
        print(f"[登录失败] {resp}")
        return False

    def _read_password(self) -> str:
        """读取密码, 兼容 Windows 各类终端"""
        try:
            return getpass.getpass("密码 (输入时不显示): ").strip()
        except Exception:
            # getpass 在某些终端(如 IDE 内嵌终端)不可用, 回退到明文输入
            print("[警告] 当前终端不支持隐藏输入, 密码将明文显示")
            return input("密码 (明文): ").strip()

    def login_qrcode(self) -> bool:
        """二维码登录"""
        print("[登录] 正在获取二维码...")
        try:
            self.api = ChaoXingAPI()
        except Exception as e:
            print(f"[错误] 初始化 API 失败: {e}")
            traceback.print_exc()
            return False

        if self.proxy_url:
            patch_session_proxy(self.api.session, self.proxy_url)

        try:
            self.api.qr_get()
            url = self.api.qr_geturl()
        except Exception as e:
            print(f"[错误] 获取二维码失败: {e}")
            traceback.print_exc()
            print("[提示] 可能是网络不通, 试试加 --proxy 10808")
            return False

        # 打印二维码链接
        print(f"\n[二维码] 请用浏览器打开以下链接, 用学习通 APP 扫码:")
        print(f"\n  {url}\n")

        # 尝试在终端画二维码
        try:
            from qrcode import QRCode
            qr = QRCode(border=1)
            qr.add_data(url)
            qr.make()
            qr.print_ascii()
            print()
        except Exception:
            pass

        print("[二维码] 等待扫描...")
        scanned = False
        while True:
            try:
                status = self.api.login_qr()
            except Exception as e:
                print(f"[错误] 二维码轮询失败: {e}")
                traceback.print_exc()
                return False

            if status.get("status") == True:
                try:
                    self.api.accinfo()
                except Exception as e:
                    print(f"[错误] 获取账号信息失败: {e}")
                    return False
                print(f"[登录成功] {self.api.acc.name} ({self.api.acc.school})")
                return True

            stype = status.get("type", "")
            if stype == "1":
                print("[二维码] 验证错误, 请重试")
                return False
            elif stype == "2":
                print("[二维码] 已失效, 请重试")
                return False
            elif stype == "4":
                if not scanned:
                    print("[二维码] 已扫描, 请在手机上确认...")
                    scanned = True
            time.sleep(1.5)

    def login(self) -> bool:
        """交互式登录"""
        print("\n" + "=" * 50)
        print("  学习通题目爬虫")
        print("=" * 50)
        while True:
            choice = input("\n选择登录方式:\n  [1] 手机号+密码\n  [2] 二维码\n  [q] 退出\n> ").strip()
            if choice == "1":
                phone = input("手机号: ").strip()
                if not phone:
                    print("[提示] 手机号不能为空")
                    continue
                passwd = _read_password("密码 (输入时不显示): ")
                if not passwd:
                    print("[提示] 密码不能为空")
                    continue
                return self.login_password(phone, passwd)
            elif choice == "2":
                return self.login_qrcode()
            elif choice.lower() == "q":
                sys.exit(0)
            else:
                print("无效选择, 请输入 1、2 或 q")

    def fetch_all_courses(self) -> list:
        """获取所有课程"""
        print("\n[拉取] 正在获取课程列表...")
        classes = self.api.fetch_classes()
        print(f"[拉取] 共 {len(classes.classes)} 门课程:")
        for i, c in enumerate(classes.classes):
            status = "已结课" if c.state else "进行中"
            print(f"  [{i}] {c.name} ({c.teacher_name}) - {status}")
        return classes

    def scrape_course(self, classes, course_index: int):
        """爬取单个课程的所有题目"""
        c = classes.classes[course_index]
        print(f"\n{'=' * 50}")
        print(f"[课程] {c.name}")
        print(f"[信息] 教师: {c.teacher_name} | 状态: {'已结课' if c.state else '进行中'}")
        print(f"{'=' * 50}")

        # 获取章节
        print("[拉取] 正在获取章节列表...")
        chap = classes.fetch_chapters_by_index(course_index)
        print(f"[拉取] 共 {len(chap.chapters)} 个章节")

        # 遍历章节提取题目
        chapters_data = []
        total_questions = 0

        for idx in range(len(chap.chapters)):
            ch = chap.chapters[idx]
            print(f"\n  [章节] {ch.label} {ch.name}")

            try:
                points = chap.fetch_points_by_index(idx)
            except Exception as e:
                self._log_error(f"    [跳过] 获取章节卡片失败: {e}", e)
                continue

            chapter_questions = []
            for point in points:
                if point.__class__.__name__ != "ChapterExam":
                    continue

                try:
                    prefetch_ok = point.pre_fetch()
                except Exception as e:
                    self._log_error(f"    [跳过] 预拉取异常: {e}", e)
                    continue
                # 标题已在 pre_fetch 中从 attachment 提取
                exam_title = getattr(point, 'title', None) or '未知'

                # pre_fetch 返回 False = 试题已提交, 跳过提交检查继续爬取
                if not prefetch_ok:
                    print(f"    [已提交] 《{exam_title}》")

                try:
                    fetch_ok = point.fetch()
                except Exception as e:
                    self._log_error(f"    [跳过] 题目拉取异常 ({exam_title}): {e}", e)
                    continue

                if not fetch_ok:
                    # 检测具体原因
                    reason = "未知原因"
                    if hasattr(point, 'title') and point.title:
                        pass  # title 被设置了说明没走无权限/已批阅分支, 是后面的解析失败
                    print(f"    [跳过] 《{exam_title}》无法拉取 (可能: 无权限需先完成前置任务/已批阅/页面结构变化)")
                    continue

                print(f"    [试题] {point.title} → {len(point.questions)} 道题")

                for q in point.questions:
                    chapter_questions.append({
                        "id": q.q_id,
                        "type": q.q_type.value,
                        "question": q.value,
                        "options": {k: v for k, v in q.answers.items()} if q.answers else {},
                        "answer": q.answer,
                    })

                time.sleep(0.5)

            if chapter_questions:
                chapters_data.append({
                    "label": ch.label,
                    "name": ch.name,
                    "questions": chapter_questions,
                })
                total_questions += len(chapter_questions)

        # 导出
        if total_questions == 0:
            print(f"\n[结果] 该课程未提取到任何题目 (可能所有试题均已批阅或无可答题试题)")
            return

        print(f"\n[结果] 共提取 {total_questions} 道题目")
        print("[导出] 正在导出 Word 文档...")
        export_docx(c.name, chapters_data, self.output_dir)

    def scrape_all(self, classes):
        """爬取所有课程"""
        for i in range(len(classes.classes)):
            self.scrape_course(classes, i)
            print()  # 课程间空行

    def run(self, course_filter: str = None, scrape_all: bool = False):
        """主入口"""
        # 1. 登录
        if not self.login():
            print("[错误] 登录失败")
            return

        # 2. 获取课程
        try:
            classes = self.fetch_all_courses()
        except Exception as e:
            print(f"[错误] 获取课程失败: {e}")
            return

        if len(classes.classes) == 0:
            print("[提示] 该账号下没有课程")
            return

        # 3. 爬取
        if scrape_all:
            self.scrape_all(classes)
        elif course_filter:
            # 按名称或序号匹配
            found = False
            # 尝试按序号
            try:
                idx = int(course_filter)
                if 0 <= idx < len(classes.classes):
                    self.scrape_course(classes, idx)
                    found = True
            except ValueError:
                pass
            # 尝试按名称匹配
            if not found:
                for i, c in enumerate(classes.classes):
                    if course_filter in c.name:
                        self.scrape_course(classes, i)
                        found = True
                        break
            if not found:
                print(f"[错误] 未找到匹配的课程: {course_filter}")
        else:
            # 交互式选择
            while True:
                choice = input("\n输入课程序号 (多个用逗号分隔, 'all' 全部, 'q' 退出): ").strip()
                if choice.lower() == "q":
                    break
                if choice.lower() == "all":
                    self.scrape_all(classes)
                    break
                for part in choice.split(","):
                    part = part.strip()
                    try:
                        idx = int(part)
                        if 0 <= idx < len(classes.classes):
                            self.scrape_course(classes, idx)
                        else:
                            print(f"[跳过] 无效序号: {idx}")
                    except ValueError:
                        print(f"[跳过] 无效输入: {part}")
                        continue

        print("\n[完成] 所有题目已导出为 Word 文档, 在 output/ 目录下")


# ============================================================
# 入口
# ============================================================

if __name__ == "__main__":
    try:
        parser = argparse.ArgumentParser(description="学习通题目爬虫")
        parser.add_argument("--proxy", type=int, default=None, help="代理端口 (如: 10808)")
        parser.add_argument("--debug", action="store_true", help="调试模式, 输出完整错误堆栈")
        parser.add_argument("--diagnose", action="store_true", help="诊断模式, 扫描章节卡片的模块类型")
        parser.add_argument("--all", action="store_true", help="爬取所有课程")
        parser.add_argument("--course", type=str, default=None, help="指定课程名称或序号")
        args = parser.parse_args()

        if args.diagnose:
            import utils
            import re as _re2
            scraper = QuestionScraper(proxy_port=args.proxy, debug=True)
            if not scraper.login():
                sys.exit(1)
            try:
                utils.save_session(scraper.api.ck_dump(), scraper.api.acc)
            except Exception:
                pass
            classes = scraper.fetch_all_courses()
            choice = input("\n选择课程序号: ").strip()
            chap = classes.fetch_chapters_by_index(int(choice))

            from cxapi import get_dc, calc_infenc
            all_modules = {}
            exam_list = []
            for idx in range(len(chap.chapters)):
                ch = chap.chapters[idx]
                params = {
                    "id": ch.chapter_id, "courseid": chap.courseid,
                    "fields": "id,parentnodeid,indexorder,label,layer,name,begintime,createtime,lastmodifytime,status,jobUnfinishedCount,clickcount,openlock,card.fields(id,knowledgeid,title,knowledgeTitile,description,cardorder).contentcard(all)",
                    "view": "json", "token": "4faa8662c59590c6f43ae9fe5b002b42",
                    "_time": str(int(time.time() * 1000)),
                }
                resp = chap.session.get(
                    "https://mooc1-api.chaoxing.com/gas/knowledge",
                    params={**params, "inf_enc": calc_infenc(params)},
                )
                cards_data = resp.json()
                for data_item in cards_data.get("data", []):
                    for card in data_item.get("card", {}).get("data", []):
                        desc = card.get("description", "")
                        if not desc:
                            continue
                        modules = _re2.findall(r'module="(\w+)"', desc)
                        for m in modules:
                            all_modules.setdefault(m, []).append({
                                "chapter": f"{ch.label} {ch.name}",
                                "title": card.get("title", ""),
                            })
                        if "work" in modules:
                            has_jobid = bool(_re2.search(r'"jobid"\s*:\s*"', desc))
                            exam_list.append({
                                "chapter": f"{ch.label} {ch.name}",
                                "title": card.get("title", ""),
                                "has_jobid": has_jobid,
                            })
                if idx % 20 == 0:
                    print(f"  扫描: {idx+1}/{len(chap.chapters)}")

            print(f"\n=== 模块类型 ===")
            for m, items in sorted(all_modules.items()):
                print(f"  [{m}] x{len(items)}")
            submitted = [e for e in exam_list if e["has_jobid"]]
            unsubmitted = [e for e in exam_list if not e["has_jobid"]]
            print(f"\n=== 试题 ===")
            print(f"  未提交: {len(unsubmitted)}  已提交(有jobid): {len(submitted)}")
            if submitted:
                print(f"\n  已提交试题:")
                for e in submitted:
                    print(f"    [{e['chapter']}] {e['title']}")
            else:
                print("  没有发现已提交试题")
            sys.exit(0)

        scraper = QuestionScraper(proxy_port=args.proxy, debug=args.debug)
        scraper.run(course_filter=args.course, scrape_all=args.all)
    except KeyboardInterrupt:
        print("\n[退出] 用户中断")
    except Exception:
        print("\n[异常] 程序出错, 详细信息如下:\n")
        traceback.print_exc()
    finally:
        # Windows 下保持窗口不闪退
        if sys.platform == "win32":
            input("\n按回车键退出...")
