"""Word export routes."""

import tempfile
import threading
import uuid
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Blueprint, flash, jsonify, redirect, request, send_file, session, url_for

from web.app import get_api, _load_api_for_phone

export_bp = Blueprint("export", __name__)

QUESTION_TYPE_NAMES = {
    0: "单选题", 1: "多选题", 2: "填空题", 3: "判断题", 4: "简答题",
    5: "名词解释", 6: "论述题", 7: "计算题", 8: "其它", 9: "分录题",
    10: "资料题", 11: "连线题", 13: "排序题", 14: "完型填空",
    15: "阅读理解", 18: "口语题", 19: "听力题",
}

# In-memory task tracker: task_id -> {current, total, done, filepath, error}
_export_tasks: dict[str, dict] = {}


def _sanitize_filename(name: str) -> str:
    forbidden = '<>:"/\\|?*'
    for ch in forbidden:
        name = name.replace(ch, "_")
    return name.strip()


def _do_export(index: int, task_id: str, phone: str):
    """Run export in background thread, updating _export_tasks."""
    task = _export_tasks[task_id]
    try:
        api = _load_api_for_phone(phone)
        if not api:
            task["error"] = "请先登录"
            task["done"] = True
            return

        classes = api.fetch_classes()
        course = classes.classes[index]
        chap = classes.fetch_chapters_by_index(index)
        total_chapters = len(chap.chapters)
        task["total"] = total_chapters

        chapters_data = []
        total_q = 0

        for ci in range(total_chapters):
            ch = chap.chapters[ci]
            try:
                points = chap.fetch_points_by_index(ci)
            except Exception:
                task["current"] = ci + 1
                continue

            chapter_questions = []
            for pt in points:
                if pt.__class__.__name__ != "ChapterExam":
                    continue
                try:
                    pt.pre_fetch()
                    pt.fetch()
                except Exception:
                    continue
                for q in getattr(pt, "questions", []):
                    chapter_questions.append({
                        "id": q.q_id,
                        "type": q.q_type.value,
                        "question": q.value,
                        "options": dict(q.answers) if q.answers else {},
                        "answer": q.answer or "",
                    })
                    total_q += 1

            if chapter_questions:
                chapters_data.append({
                    "label": ch.label,
                    "name": ch.name,
                    "questions": chapter_questions,
                })

            task["current"] = ci + 1

        if total_q == 0:
            task["error"] = "该课程没有可导出的题目"
            task["done"] = True
            return

        # Generate .docx
        doc = Document()
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        title_para = doc.add_heading(course.name, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(
            f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}    "
            f"章节数: {len(chapters_data)}    题目总数: {total_q}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)
        doc.add_paragraph()

        q_idx = 0
        for ch in chapters_data:
            doc.add_heading(f"{ch['label']} {ch['name']}", level=1)
            for q in ch["questions"]:
                q_idx += 1
                qtype = QUESTION_TYPE_NAMES.get(q["type"], f"未知({q['type']})")
                q_para = doc.add_paragraph()
                q_para.paragraph_format.space_after = Pt(4)
                q_run = q_para.add_run(f"{q_idx}. ({qtype}) ")
                q_run.bold = True
                q_run.font.size = Pt(11)
                q_para.add_run(q["question"]).font.size = Pt(11)

                if q["options"]:
                    for opt_label, opt_text in q["options"].items():
                        opt_para = doc.add_paragraph()
                        opt_para.paragraph_format.left_indent = Cm(1)
                        opt_para.paragraph_format.space_after = Pt(2)
                        opt_run = opt_para.add_run(f"{opt_label}. {opt_text}")
                        opt_run.font.size = Pt(10.5)

                answer = q.get("answer", "")
                if answer:
                    options = q.get("options", {})
                    answer_texts = []
                    for ch_a in answer:
                        if ch_a in options and options[ch_a]:
                            answer_texts.append(f"{ch_a}. {options[ch_a]}")
                        else:
                            answer_texts.append(ch_a)
                    full_answer = "\n".join(answer_texts) if len(answer_texts) > 1 else answer_texts[0]
                    ans_para = doc.add_paragraph()
                    ans_para.paragraph_format.space_before = Pt(6)
                    ans_para.paragraph_format.space_after = Pt(8)
                    ans_run = ans_para.add_run(f"✓ 正确答案: {full_answer}")
                    ans_run.bold = True
                    ans_run.font.size = Pt(10.5)
                    ans_run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    doc.add_paragraph()

        safe_name = _sanitize_filename(course.name)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()

        task["filepath"] = tmp.name
        task["filename"] = f"{safe_name}.docx"
        task["chapter_count"] = len(chapters_data)
        task["question_count"] = total_q
        task["done"] = True
    except Exception as e:
        task["error"] = str(e)
        task["done"] = True


@export_bp.route("/api/export/<int:index>/start", methods=["POST"])
def export_start(index: int):
    """Start an export task, return task_id for polling."""
    phone = session.get("phone")
    if not phone:
        return jsonify({"error": "请先登录"}), 401

    task_id = uuid.uuid4().hex[:12]
    _export_tasks[task_id] = {"current": 0, "total": 0, "done": False}

    t = threading.Thread(target=_do_export, args=(index, task_id, phone), daemon=True)
    t.start()

    return jsonify({"task_id": task_id})


@export_bp.route("/api/export/status/<task_id>")
def export_status(task_id: str):
    """Poll export task progress."""
    task = _export_tasks.get(task_id)
    if not task:
        return jsonify({"done": True, "error": "任务不存在"})
    return jsonify({
        "current": task.get("current", 0),
        "total": task.get("total", 0),
        "done": task.get("done", False),
        "error": task.get("error"),
    })


@export_bp.route("/api/export/download/<task_id>")
def export_download(task_id: str):
    """Download completed export file."""
    task = _export_tasks.get(task_id)
    if not task or not task.get("filepath"):
        flash("文件不存在或已过期", "error")
        return redirect(url_for("courses.list_courses"))

    filepath = task["filepath"]
    filename = task.get("filename", "export.docx")
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
